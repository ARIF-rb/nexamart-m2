"""
Build the NexaMart M2 team-coordination artifacts:
  - tracking/nexamart_m2_tracker.xlsx   (7 tabs + hidden _lists; dropdowns; RAG; rollups)
  - docs/team_assembly_M2.docx          (report-assembly skeleton -> Google Doc)
  - docs/team_assembly_M2.md            (markdown mirror)

Run from the repo root:  python tracking/build_tracker.py
Idempotent: overwrites the three outputs each run. No external answer-key data is embedded
(before/after counts stay in the gitignored .private/ — the shared tracker is count-free).
"""
import os, re, glob

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.formatting.rule import CellIsRule
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # repo root
TASKS_DIR = os.path.join(ROOT, "docs", "tasks")

# ---------------------------------------------------------------------------
# Controlled vocab
# ---------------------------------------------------------------------------
STATUS = ["Not started", "In progress", "Blocked", "Done"]
RESULT = ["Pending", "Pass", "Fail"]
YESNO = ["No", "Yes"]

# ---------------------------------------------------------------------------
# Static data (public fields only)
# ---------------------------------------------------------------------------
# ID, Category, Title, Owner, Source Silver, Resolution method/b_classification
ANOMALIES = [
    ("A1", "A", "Cancelled EC order still in revenue", "M5", "silver_ec_orders", "Zero cancelled revenue"),
    ("A2", "A", "Payment captured after cancellation", "M5", "silver_pg_transactions", "Flag reversal-required"),
    ("A3", "A", "Tax/shipping inclusion mismatch", "M5", "silver_ec_orders; silver_pos_transactions", "Normalise to tax-exclusive"),
    ("A4", "A", "NexaLocal seller-marked-sold as confirmed revenue", "M6", "silver_nl_listing_events", "Relabel ESTIMATED (feeds B6)"),
    ("A5", "A", "Website ATP>0 while warehouse physical=0", "Lead", "silver_wh_inventory_snapshots", "Correct ATP to 0"),
    ("A6", "A", "Negative ATP / sellable qty", "M4", "silver_si_inventory_snapshots", "Correct to 0 + oversell flag"),
    ("A7", "A", "Returned stock sellable before inspection", "Lead", "silver_rr_return_receipts", "Zero pre-inspection restock"),
    ("A8", "A", "Missing snapshot days (1-7 Aug ramp-up)", "Lead", "silver_si_inventory_snapshots", "Reconstruct -> RECONSTRUCTED/INFERRED"),
    ("A9", "A", "Same SKU -> different products", "M3", "silver_product_master", "Canonical product (catalogue wins)"),
    ("A10", "A", "Open-box restocked as NEW", "M4", "silver_rr_return_receipts", "Correct condition to open-box"),
    ("A11", "A", "Placeholder customer 9999 collision", "M2", "silver_customer_master; silver_ec_orders", "Rekey guests to GUEST-{session}"),
    ("A12", "A", "Listing marked sold then relisted", "Lead", "silver_nl_listings", "Link + exclude from GMV; reliability=LOW"),
    ("A13", "A", "Coordinated fake listing ring", "M6", "silver_nl_listings", "Flag risk tier; exclude"),
    ("A14", "A", "Delivered before shipped", "Lead", "silver_dc_delivery_events", "+36h median; >72h manual review"),
    ("A15", "A", "Review before delivery", "M6", "silver_rv_reviews", "verified_purchase=FALSE"),
    ("A16", "A", "Duplicate complaint cases", "M6", "silver_cs_cases", "Dedupe to canonical case key"),
    ("B1", "B", "Campaign-window attribution", "Lead", "silver_ec_orders; silver_ws_sessions", "Attribute (b_class=ATTRIBUTED, conf 0.85)"),
    ("B2", "B", "Partial refund period attribution", "Lead", "silver_rr_refund_events", "Recognise in return period"),
    ("B3", "B", "Inventory movement without ref order", "Lead", "silver_wh_inventory_movements", "Classify PROBABLE_MISSING_REF (INFERRED)"),
    ("B4", "B", "NexaLocal listing -> catalogue match", "M3", "silver_nl_listings", "Match >=0.75 / 0.65-0.75 review / <0.65 unmatched"),
    ("B5", "B", "Cross-channel identity resolution", "M2", "silver_customer_master", "Probabilistic merge >=0.90"),
    ("B6", "B", "Estimated Classified GMV model", "M2", "silver_nl_listing_events", "Weighted signals + -/+35% band, ESTIMATED"),
    ("B7", "B", "BOPIS Completed without pickup event", "Lead", "silver_ec_orders; silver_dc_delivery_events", "Treat as fulfilled; collection_unconfirmed"),
    ("B8", "B", "Seller trust composite score", "M2", "silver_ts_sellers", "Weighted composite -> 5 risk tiers"),
]

VALIDATION = [
    (1, "Completeness / row counts", "Lead"),
    (2, "Referential Integrity (no orphan FKs)", "Lead"),
    (3, "Grain Violations", "M3"),
    (4, "Additive Fact Sanity (net=gross-disc-return)", "M5 / M4"),
    (5, "Semi-Additive Guard (no SUM(ATP) across dates)", "Lead"),
    (6, "Metric Certainty Completeness", "M2"),
    (7, "Campaign Period Coverage (>=1 row/fact)", "Lead"),
    (8, "Inventory Balance Reconciliation", "M4"),
    (9, "Classified Certainty Segregation", "M2"),
    (10, "Temporal Consistency (delivered-after-shipped)", "Lead"),
]

DELIVERABLES = [
    ("Written report (PDF)", "report/nexamart_m2_report.pdf", "Lead"),
    ("Anomaly resolution notebook", "notebooks/05_anomaly_resolution.ipynb", "All (resolution owners)"),
    ("Gold rebuild notebook", "notebooks/06_gold_rebuild.ipynb", "Lead + fact owners"),
    ("Anomaly discovery SQL", "sql/anomaly_discovery.sql", "All"),
    ("Anomaly resolution SQL", "sql/anomaly_resolution.sql", "All"),
    ("Validation suite SQL", "sql/validation_suite.sql", "Lead"),
    ("KPI views SQL", "sql/kpi_views.sql", "All (KPI owners)"),
    ("Dashboard file (.pbix/.twbx)", "dashboard/nexamart_dashboard.*", "All (page owners)"),
    ("Dashboard screenshots", "dashboard/screenshots/", "All"),
    ("Presentation (5-10 slides)", "presentation/nexamart_m2_presentation.*", "Lead"),
    ("Submission ZIP", "nexamart_m2_group_[N].zip", "Lead"),
]

# Anomaly ownership grouped for the assembly doc S1 markers (preserve order)
S1_OWNER_ORDER = ["M5", "M6", "Lead", "M4", "M3", "M2"]


def s1_groups():
    groups = {}
    for aid, cat, title, owner, *_ in ANOMALIES:
        groups.setdefault(owner, []).append(aid)
    return [(o, groups[o]) for o in S1_OWNER_ORDER if o in groups]


# ---------------------------------------------------------------------------
# Parse tasks from docs/tasks/*.md
# ---------------------------------------------------------------------------
def member_name(fname):
    base = os.path.splitext(os.path.basename(fname))[0]
    return "Lead" if base.lower() == "lead" else base  # M2..M6


def parse_tasks():
    rows = []
    files = ["lead.md", "M2.md", "M3.md", "M4.md", "M5.md", "M6.md"]
    day_re = re.compile(r"^##\s+Day\s+(\d+)")
    task_re = re.compile(r"^###\s+Task\s+(\d+)\s+[—-]\s+(.*)$")
    est_re = re.compile(r"\*\*Est:\*\*\s*([\d.]+)\s*hr", re.I)
    blk_re = re.compile(r"\*\*Blocked by:\*\*\s*(.*)$", re.I)
    for f in files:
        path = os.path.join(TASKS_DIR, f)
        member = member_name(path)
        day = ""
        with open(path, encoding="utf-8") as fh:
            lines = fh.read().splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            dm = day_re.match(line)
            if dm:
                day = int(dm.group(1))
                i += 1
                continue
            tm = task_re.match(line)
            if tm:
                tnum = int(tm.group(1))
                title = tm.group(2).strip()
                est, blocked = "", ""
                j = i + 1
                while j < len(lines) and not task_re.match(lines[j]) and not day_re.match(lines[j]):
                    em = est_re.search(lines[j])
                    if em and est == "":
                        est = float(em.group(1))
                    bm = blk_re.search(lines[j])
                    if bm and blocked == "":
                        blocked = bm.group(1).strip()
                    j += 1
                rows.append([member, day, tnum, title, est, blocked])
                i = j
                continue
            i += 1
    return rows


# ---------------------------------------------------------------------------
# Parse KPIs from docs/kpi_register.md
# ---------------------------------------------------------------------------
def parse_kpis():
    rows = []
    path = os.path.join(ROOT, "docs", "kpi_register.md")
    with open(path, encoding="utf-8") as fh:
        for line in fh:
            if not line.strip().startswith("|"):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if len(cells) < 7 or not cells[0].isdigit():
                continue
            # # | KPI | View | Group | Certainty | Owner | Source
            num, kpi, view, group, certainty, owner = cells[0], cells[1], cells[2], cells[3], cells[4], cells[5]
            certainty = certainty.replace("**", "")
            rows.append([kpi, view, group, owner, certainty])
    return rows


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------
HDR_FILL = PatternFill("solid", fgColor="1F3864")
HDR_FONT = Font(bold=True, color="FFFFFF", size=11)
TITLE_FONT = Font(bold=True, size=14, color="1F3864")
THIN = Side(style="thin", color="BFBFBF")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
RAG = {
    "Done": ("C6EFCE", "006100"),
    "In progress": ("FFEB9C", "9C6500"),
    "Blocked": ("FFC7CE", "9C0006"),
    "Not started": ("D9D9D9", "3F3F3F"),
    "Pass": ("C6EFCE", "006100"),
    "Fail": ("FFC7CE", "9C0006"),
    "Pending": ("D9D9D9", "3F3F3F"),
}


def style_header(ws, ncols, row=1):
    for c in range(1, ncols + 1):
        cell = ws.cell(row=row, column=c)
        cell.fill = HDR_FILL
        cell.font = HDR_FONT
        cell.alignment = Alignment(vertical="center", wrap_text=True)
        cell.border = BORDER


def add_table(ws, headers, data, widths=None):
    ws.append(headers)
    style_header(ws, len(headers))
    for r in data:
        ws.append(r)
    ws.freeze_panes = "A2"
    if widths:
        for i, w in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = w
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=len(headers)):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = BORDER


def apply_status_dv(ws, col_letter, first, last, list_range, values):
    dv = DataValidation(type="list", formula1=f"={list_range}", allow_blank=True)
    ws.add_data_validation(dv)
    dv.add(f"{col_letter}{first}:{col_letter}{last}")
    rng = f"{col_letter}{first}:{col_letter}{last}"
    for v in values:
        fill = PatternFill("solid", fgColor=RAG[v][0])
        font = Font(color=RAG[v][1])
        ws.conditional_formatting.add(
            rng, CellIsRule(operator="equal", formula=[f'"{v}"'], fill=fill, font=font))


def build_xlsx(path, tasks, kpis):
    wb = Workbook()

    # hidden _lists
    lists = wb.active
    lists.title = "_lists"
    for i, v in enumerate(STATUS, 1):
        lists.cell(row=i, column=1, value=v)
    for i, v in enumerate(RESULT, 1):
        lists.cell(row=i, column=2, value=v)
    for i, v in enumerate(YESNO, 1):
        lists.cell(row=i, column=3, value=v)
    lists.sheet_state = "hidden"
    L_STATUS = "_lists!$A$1:$A$4"
    L_RESULT = "_lists!$B$1:$B$3"
    L_YESNO = "_lists!$C$1:$C$2"

    # 2. Anomalies
    an = wb.create_sheet("Anomalies")
    add_table(an,
        ["ID", "Cat", "Title", "Owner", "Source Silver", "Resolution method / b_class",
         "Detection", "Resolved", "Gold rebuilt", "Re-verified->0", "Status", "Notes"],
        [[a[0], a[1], a[2], a[3], a[4], a[5], "No", "No", "No", "No", "Not started", ""] for a in ANOMALIES],
        widths=[7, 5, 38, 7, 30, 34, 11, 10, 12, 14, 13, 24])
    last = 1 + len(ANOMALIES)
    for col in ("G", "H", "I", "J"):
        apply_status_dv(an, col, 2, last, L_YESNO, [])
    apply_status_dv(an, "K", 2, last, L_STATUS, STATUS)

    # 3. Tasks
    tk = wb.create_sheet("Tasks")
    add_table(tk,
        ["Member", "Day", "Task#", "Title", "Est hrs", "Blocked by", "Status", "Notes"],
        [[t[0], t[1], t[2], t[3], (t[4] if t[4] != "" else None), t[5], "Not started", ""] for t in tasks],
        widths=[9, 6, 7, 52, 9, 30, 13, 22])
    tlast = 1 + len(tasks)
    apply_status_dv(tk, "G", 2, tlast, L_STATUS, STATUS)

    # 4. KPI Views
    kp = wb.create_sheet("KPI Views")
    add_table(kp,
        ["KPI", "View", "Group", "Owner", "Certainty", "View created", "Validated", "Status"],
        [[k[0], k[1], k[2], k[3], k[4], "No", "No", "Not started"] for k in kpis],
        widths=[34, 36, 18, 8, 12, 13, 11, 13])
    klast = 1 + len(kpis)
    apply_status_dv(kp, "F", 2, klast, L_YESNO, [])
    apply_status_dv(kp, "G", 2, klast, L_YESNO, [])
    apply_status_dv(kp, "H", 2, klast, L_STATUS, STATUS)

    # 5. Validation
    vd = wb.create_sheet("Validation")
    add_table(vd,
        ["Check#", "Name", "Owner", "Iter-1 result", "Iter-2 result", "Final result", "Notes"],
        [[c[0], c[1], c[2], "", "", "Pending", ""] for c in VALIDATION],
        widths=[8, 46, 10, 14, 14, 13, 26])
    vlast = 1 + len(VALIDATION)
    apply_status_dv(vd, "F", 2, vlast, L_RESULT, RESULT)

    # 6. Deliverables
    dl = wb.create_sheet("Deliverables")
    add_table(dl,
        ["Artifact", "Path", "Owner", "Status"],
        [[d[0], d[1], d[2], "Not started"] for d in DELIVERABLES],
        widths=[34, 42, 24, 13])
    dlast = 1 + len(DELIVERABLES)
    apply_status_dv(dl, "D", 2, dlast, L_STATUS, STATUS)

    # 7. Standup log
    su = wb.create_sheet("Standup log")
    add_table(su, ["Date", "Member", "Yesterday", "Today", "Blockers"],
              [["", "", "", "", ""] for _ in range(40)],
              widths=[14, 10, 40, 40, 30])
    apply_status_dv(su, "B", 2, 41, "_lists!$D$1:$D$7", [])
    for i, m in enumerate(["Lead", "M2", "M3", "M4", "M5", "M6"], 1):
        lists.cell(row=i, column=4, value=m)

    # 1. Overview (first visible tab)
    ov = wb.create_sheet("Overview", 0)
    ov["A1"] = "NexaMart M2 — Group Execution Tracker"
    ov["A1"].font = TITLE_FONT
    meta = [
        ("Milestone", "M2 of 2 — Validate, Fix, Analyse, Conclude"),
        ("Submission", "Single ZIP: nexamart_m2_group_[N].zip"),
        ("Campaign window", "8-28 Aug 2024 (baseline 1 Mar-31 Jul; ramp 1-7 Aug; post 29 Aug-14 Sep)"),
        ("Repo / branch", "github.com/ARIF-rb/nexamart-m2 — working branch: master"),
        ("Deadline", "Day 10 EOD"),
    ]
    r = 3
    for k, v in meta:
        ov.cell(row=r, column=1, value=k).font = Font(bold=True)
        ov.cell(row=r, column=2, value=v)
        r += 1

    r += 1
    ov.cell(row=r, column=1, value="PROGRESS").font = Font(bold=True, color="1F3864")
    r += 1
    rollups = [
        ("Anomalies resolved", '=COUNTIF(Anomalies!K2:K{},"Done")&"/"&{}'.format(last, len(ANOMALIES))),
        ("KPI views done", '=COUNTIF(\'KPI Views\'!H2:H{},"Done")&"/"&{}'.format(klast, len(kpis))),
        ("Validation checks passing", '=COUNTIF(Validation!F2:F{},"Pass")&"/"&{}'.format(vlast, len(VALIDATION))),
        ("Deliverables done", '=COUNTIF(Deliverables!D2:D{},"Done")&"/"&{}'.format(dlast, len(DELIVERABLES))),
        ("Tasks done", '=COUNTIF(Tasks!G2:G{},"Done")&"/"&{}'.format(tlast, len(tasks))),
    ]
    for k, formula in rollups:
        ov.cell(row=r, column=1, value=k).font = Font(bold=True)
        ov.cell(row=r, column=2, value=formula)
        r += 1

    ov.column_dimensions["A"].width = 24
    ov.column_dimensions["B"].width = 60
    ov.column_dimensions["C"].width = 22

    wb.save(path)
    return {"Anomalies": len(ANOMALIES), "Tasks": len(tasks), "KPI Views": len(kpis),
            "Validation": len(VALIDATION), "Deliverables": len(DELIVERABLES)}


# ---------------------------------------------------------------------------
# DOCX + MD assembly skeleton
# ---------------------------------------------------------------------------
SECTIONS = [
    ("Executive Summary", "Lead", "<=1 page. Campaign question, headline reconciled numbers (GSV->NCR, Confirmed GMV, Estimated Classified GMV with band, Net Margin), one-line verdict.", []),
    ("Section 1 — Anomaly Resolution Report", None,
     "Per Category A: detection / root cause / PySpark fix / before->after / business impact. Per Category B: ambiguity / chosen interpretation / implementation / business defence + quantified alternative.",
     s1_groups()),
    ("Section 2 — Gold Rebuild Summary", "Lead",
     "Single Lead-owned block (with an embedded per-fact owner note). Which Gold tables were affected, which notebook sections re-run, row counts before vs after, which tables were rebuild-exempt.", []),
    ("Section 3 — Validation Outcomes", "Lead",
     "All 10 checks across all iterations; how many iterations; what failed and was fixed each iteration; final state.", []),
    ("Section 4 — KPI Reconciliation Report", "Lead",
     "Why all seven teams reported different numbers (Sales/Finance/Inventory/Ecommerce/Store Ops/Marketplace/Support); GSV->NCR waterfall naming + quantifying each deduction.", []),
    ("Section 5 — Campaign Performance Conclusion", "Lead",
     "Was the campaign successful? NCR, Net Margin, Inventory, Fulfilment, Customer satisfaction, Classified engagement — each with its certainty level, never conflated. A direct, defensible position.", []),
]


def marker_text(owner):
    return f"▼ {owner} — PASTE YOUR PART HERE ▼"


def end_text(owner):
    return f"▲ END {owner} ▲"


def build_docx(path):
    doc = Document()
    doc.add_heading("NexaMart M2 — Report Team Assembly", level=0)
    p = doc.add_paragraph()
    p.add_run("How to use: ").bold = True
    p.add_run("Find your member code with Ctrl+F (e.g. \"M3\"), then paste your written part "
              "BETWEEN your ▼ PASTE ▼ and ▲ END ▲ markers. Don't edit other members' blocks. "
              "Section ownership matches report/nexamart_m2_report.md. Keep certainty levels labelled "
              "(CONFIRMED / INFERRED / ESTIMATED) and never mix them in one figure. Track your progress "
              "in the Group Execution Tracker (Sheet) — tick your tasks, anomalies, and KPI views as you complete them.")
    for title, owner, guide, groups in SECTIONS:
        doc.add_heading(title, level=1)
        gi = doc.add_paragraph(guide); gi.runs[0].italic = True
        if groups:  # S1 — per-anomaly owners
            for ow, ids in groups:
                doc.add_heading(f"{ow} — anomalies {', '.join(ids)}", level=2)
                _marker_block(doc, ow)
        else:
            if owner == "Lead" and title.startswith("Section 2"):
                note = doc.add_paragraph("Per-fact contributors feed Lead here: M5 fact_ecommerce_order_line/"
                                         "store-sale/return; M4 warehouse snapshot/dim_store; M3 dim_product; "
                                         "M2 dim_customer/identity-bridge/seller-risk-tier; M6 review/listing facts.")
                note.runs[0].italic = True
            _marker_block(doc, owner)
    doc.save(path)


def _marker_block(doc, owner):
    m = doc.add_paragraph(); r = m.add_run(marker_text(owner)); r.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    doc.add_paragraph("")  # paste area
    e = doc.add_paragraph(); er = e.add_run(end_text(owner)); er.bold = True; er.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)


def build_md(path):
    lines = ["# NexaMart M2 — Report Team Assembly", "",
             "**How to use:** Find your member code with Ctrl+F (e.g. `M3`), paste your part between your "
             "`▼ PASTE ▼` and `▲ END ▲` markers. Section ownership matches "
             "`report/nexamart_m2_report.md`. Keep certainty levels labelled and never mix them in one figure. "
             "Track your progress in the Group Execution Tracker (Sheet) — tick your tasks, anomalies, and KPI views as you complete them.", ""]
    for title, owner, guide, groups in SECTIONS:
        lines += [f"## {title}", "", f"*{guide}*", ""]
        if groups:
            for ow, ids in groups:
                lines += [f"### {ow} — anomalies {', '.join(ids)}", "",
                          f"`{marker_text(ow)}`", "", "", f"`{end_text(ow)}`", ""]
        else:
            if owner == "Lead" and title.startswith("Section 2"):
                lines += ["*Per-fact contributors feed Lead here: M5 fact_ecommerce_order_line/store-sale/return; "
                          "M4 warehouse snapshot/dim_store; M3 dim_product; M2 dim_customer/identity-bridge/"
                          "seller-risk-tier; M6 review/listing facts.*", ""]
            lines += [f"`{marker_text(owner)}`", "", "", f"`{end_text(owner)}`", ""]
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines) + "\n")


# ---------------------------------------------------------------------------
if __name__ == "__main__":
    tasks = parse_tasks()
    kpis = parse_kpis()
    xlsx = os.path.join(ROOT, "tracking", "nexamart_m2_tracker.xlsx")
    docx_path = os.path.join(ROOT, "docs", "team_assembly_M2.docx")
    md_path = os.path.join(ROOT, "docs", "team_assembly_M2.md")
    counts = build_xlsx(xlsx, tasks, kpis)
    build_docx(docx_path)
    build_md(md_path)
    print("tasks parsed:", len(tasks), "| per-member:",
          {m: sum(1 for t in tasks if t[0] == m) for m in ["Lead", "M2", "M3", "M4", "M5", "M6"]})
    print("kpis parsed:", len(kpis))
    print("xlsx data-row counts:", counts)
    print("wrote:", xlsx)
    print("wrote:", docx_path)
    print("wrote:", md_path)
